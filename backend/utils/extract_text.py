import os
from docx import Document
import pdfplumber
import logging
import re

logging.basicConfig(level=logging.ERROR) 

def extract_text_from_pdf(pdf_path):
    """
    Extracts text from a PDF file using pdfplumber.

    Args:
        pdf_path (str): The file path of the PDF.

    Returns:
        str: Extracted text from the PDF or None if an error occurs.
    """
    try:
        with pdfplumber.open(pdf_path) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() + "\n" if page.extract_text() else ""
        print("Getting resume text from pdf...")
        return text
    except Exception as e:
        logging.error(f"Error extracting PDF: {e}")
        return None

def extract_text_from_docx(docx_path):
    """
    Extracts text from a DOCX file, including both paragraphs and table content.

    Args:
        docx_path (str): The file path of the DOCX document.

    Returns:
        str: Extracted text from the DOCX file or None if an error occurs.
    """
    try:
        doc = Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        print("Getting resume text from doc...")
        return text if text else "No text found in document"
    except Exception as e:
        logging.error(f"Error extracting DOCX: {e}")
        return None

def extract_text(file_path):
    """
    Extracts text from a PDF or DOCX file.
    Redacts contact information (phone numbers, emails, LinkedIn/GitHub links).
    
    Args:
        file_path (str): The path of the file to be processed.

    Returns:
        dict: A dictionary containing the extracted resume text.
              Example: {"resume_text": "Extracted text here"}
    """
    ext = os.path.splitext(file_path)[-1].lower()
    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif ext == ".docx":
        text = extract_text_from_docx(file_path)
    else:
        return "Unsupported file format"

    if text is None: 
        return "Error during text extraction"
    
    phone_pattern = r"(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}"
    text = re.sub(phone_pattern, "[PHONE REDACTED]", text)

    # Redact email addresses
    email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
    text = re.sub(email_pattern, "[EMAIL REDACTED]", text)

    # Redact LinkedIn links 
    linkedin_pattern = r"linkedin\.com\/in\/[a-zA-Z0-9_-]+"
    text = re.sub(linkedin_pattern, "[LINKEDIN REDACTED]", text)

    # Redact GitHub links 
    github_pattern = r"github\.com\/[a-zA-Z0-9_-]+"
    text = re.sub(github_pattern, "[GITHUB REDACTED]", text)

    print("Processed Resume text...")
    return {"resume_text": text}  

